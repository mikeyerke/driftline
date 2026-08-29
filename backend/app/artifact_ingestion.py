"""Ephemeral, bounded text extraction for PM decision artifacts."""

from __future__ import annotations

import base64
import binascii
import io
import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import PurePath

from docx import Document
from pypdf import PdfReader

MAX_ARTIFACT_BYTES = 4 * 1024 * 1024
MAX_ARTIFACT_TEXT_CHARS = 20_000
MAX_PDF_PAGES = 40
MAX_DOCX_ENTRIES = 1_000
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024

_TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".json"}
_PDF_EXTENSIONS = {".pdf"}
_DOCX_EXTENSIONS = {".docx"}
SUPPORTED_ARTIFACT_EXTENSIONS = _TEXT_EXTENSIONS | _PDF_EXTENSIONS | _DOCX_EXTENSIONS


class ArtifactIngestionError(ValueError):
    """Raised when an uploaded artifact cannot be safely read."""


@dataclass(frozen=True)
class ExtractedArtifact:
    text: str
    filename: str
    file_type: str
    byte_count: int
    truncated: bool
    redactions_applied: int


def _safe_filename(filename: str) -> str:
    name = PurePath(filename.strip()).name
    if not name or len(name) > 180 or any(ord(character) < 32 for character in name):
        raise ArtifactIngestionError("Artifact filename is invalid")
    return name


def _decode_base64(content_base64: str) -> bytes:
    if len(content_base64) > ((MAX_ARTIFACT_BYTES + 2) // 3) * 4 + 16:
        raise ArtifactIngestionError("Artifact exceeds the 4 MB upload limit")
    try:
        payload = base64.b64decode(content_base64, validate=True)
    except (binascii.Error, ValueError) as exc:
        raise ArtifactIngestionError("Artifact content is not valid base64") from exc
    if not payload:
        raise ArtifactIngestionError("Artifact is empty")
    if len(payload) > MAX_ARTIFACT_BYTES:
        raise ArtifactIngestionError("Artifact exceeds the 4 MB upload limit")
    return payload


def _extract_text_file(payload: bytes, extension: str) -> str:
    try:
        text = payload.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ArtifactIngestionError("Text artifacts must use UTF-8 encoding") from exc
    if extension == ".json":
        try:
            parsed = json.loads(text)
        except json.JSONDecodeError as exc:
            raise ArtifactIngestionError("JSON artifact is malformed") from exc
        text = json.dumps(parsed, ensure_ascii=False, indent=2)
    return text


def _extract_pdf(payload: bytes) -> str:
    if not payload.startswith(b"%PDF-"):
        raise ArtifactIngestionError("Uploaded PDF has an invalid signature")
    try:
        reader = PdfReader(io.BytesIO(payload), strict=False)
        if reader.is_encrypted:
            raise ArtifactIngestionError("Encrypted PDFs are not supported")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ArtifactIngestionError(
                f"PDF exceeds the {MAX_PDF_PAGES}-page extraction limit"
            )
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages)
    except ArtifactIngestionError:
        raise
    except Exception as exc:
        raise ArtifactIngestionError("PDF text could not be extracted") from exc


def _validate_docx_archive(payload: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise ArtifactIngestionError("DOCX archive contains too many entries")
            if sum(entry.file_size for entry in entries) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ArtifactIngestionError("DOCX expands beyond the safe extraction limit")
            if "word/document.xml" not in {entry.filename for entry in entries}:
                raise ArtifactIngestionError("DOCX is missing its document body")
    except ArtifactIngestionError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise ArtifactIngestionError("DOCX archive is invalid") from exc


def _extract_docx(payload: bytes) -> str:
    _validate_docx_archive(payload)
    try:
        document = Document(io.BytesIO(payload))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(block for block in blocks if block)
    except Exception as exc:
        raise ArtifactIngestionError("DOCX text could not be extracted") from exc


def _redact_sensitive_patterns(text: str) -> tuple[str, int]:
    patterns = [
        re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.IGNORECASE),
        re.compile(r"\b(?:Bearer\s+)?(?:sk-|ghp_|github_pat_|xox[baprs]-|AIza)[A-Za-z0-9._-]{12,}\b"),
        re.compile(r"\b(?:api[_ -]?key|access[_ -]?token|client[_ -]?secret)\s*[:=]\s*[^\s,;]{8,}", re.IGNORECASE),
        re.compile(r"(?<!\d)(?:\+?1[-. (]*)?\d{3}[-. )]+\d{3}[-. ]+\d{4}(?!\d)"),
    ]
    redactions = 0
    for pattern in patterns:
        text, count = pattern.subn("[redacted sensitive value]", text)
        redactions += count
    return text, redactions


def _normalize_text(text: str) -> tuple[str, bool, int]:
    normalized = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized).strip()
    normalized, redactions = _redact_sensitive_patterns(normalized)
    if len(normalized) < 40:
        raise ArtifactIngestionError(
            "Artifact does not contain enough readable decision context"
        )
    truncated = len(normalized) > MAX_ARTIFACT_TEXT_CHARS
    return normalized[:MAX_ARTIFACT_TEXT_CHARS], truncated, redactions


def extract_uploaded_artifact(
    *, filename: str, content_base64: str
) -> ExtractedArtifact:
    """Extract bounded text without persisting the upload or its raw contents."""
    safe_name = _safe_filename(filename)
    extension = PurePath(safe_name).suffix.casefold()
    if extension not in SUPPORTED_ARTIFACT_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_ARTIFACT_EXTENSIONS))
        raise ArtifactIngestionError(f"Unsupported artifact type; use {supported}")
    payload = _decode_base64(content_base64)
    if extension in _TEXT_EXTENSIONS:
        raw_text = _extract_text_file(payload, extension)
        file_type = "text"
    elif extension in _PDF_EXTENSIONS:
        raw_text = _extract_pdf(payload)
        file_type = "pdf"
    else:
        raw_text = _extract_docx(payload)
        file_type = "docx"
    text, truncated, redactions = _normalize_text(raw_text)
    return ExtractedArtifact(
        text=text,
        filename=safe_name,
        file_type=file_type,
        byte_count=len(payload),
        truncated=truncated,
        redactions_applied=redactions,
    )


def normalize_pasted_artifact(text: str) -> ExtractedArtifact:
    """Apply the same bounds to pasted text without pretending it is a file."""
    normalized, truncated, redactions = _normalize_text(text)
    return ExtractedArtifact(
        text=normalized,
        filename="pasted-redacted-text",
        file_type="text",
        byte_count=len(text.encode("utf-8")),
        truncated=truncated,
        redactions_applied=redactions,
    )
